#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Compare formatting between template docx and generated docx.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import copy

TEMPLATE_PATH = "/Users/wen/Desktop/未命名文件夹/DeepLearning/论文/软件学报2016年排版样例模版.docx"
GENERATED_PATH = "/Users/wen/Desktop/未命名文件夹/DeepLearning/motor_osdi24_zh.docx"

STYLES_TO_CHECK = [
    "Subtitle", "作者", "Name", "摘要", "Abstract",
    "关键词", "Key words", "Heading 1", "Heading 2",
    "Body Text", "Text of Reference", "Reference"
]

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    None: "None",
}

def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 3)

def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / 12700, 2)

def pt_val(pt_obj):
    if pt_obj is None:
        return None
    return round(float(pt_obj), 2)

def color_str(color):
    if color is None:
        return "None"
    try:
        if color.type is None:
            return "None"
        return f"#{color.rgb}" if color.rgb else str(color)
    except Exception:
        return "None"

def get_line_spacing(pf):
    """Return line spacing info as string."""
    ls = pf.line_spacing
    lsr = pf.line_spacing_rule
    if ls is None:
        return "None"
    # If lsr is EXACTLY or AT_LEAST, ls is in EMU; otherwise it's a float multiplier
    from docx.enum.text import WD_LINE_SPACING
    if lsr in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
        return f"{emu_to_pt(ls)}pt ({lsr})"
    elif lsr == WD_LINE_SPACING.MULTIPLE:
        return f"{round(float(ls), 3)}x (MULTIPLE)"
    else:
        return f"{ls} (rule={lsr})"

def get_section_info(doc):
    """Extract page/section settings."""
    results = {}
    for i, section in enumerate(doc.sections):
        key = f"section[{i}]"
        results[key] = {
            "page_width_cm": emu_to_cm(section.page_width),
            "page_height_cm": emu_to_cm(section.page_height),
            "left_margin_cm": emu_to_cm(section.left_margin),
            "right_margin_cm": emu_to_cm(section.right_margin),
            "top_margin_cm": emu_to_cm(section.top_margin),
            "bottom_margin_cm": emu_to_cm(section.bottom_margin),
            "orientation": str(section.orientation),
        }
        # Check columns via XML
        sectPr = section._sectPr
        cols_el = sectPr.find(qn('w:cols'))
        if cols_el is not None:
            num = cols_el.get(qn('w:num'))
            space = cols_el.get(qn('w:space'))
            results[key]["cols_num"] = num
            results[key]["cols_space_pt"] = round(int(space) / 20, 2) if space else None
        else:
            results[key]["cols_num"] = "1 (default)"
            results[key]["cols_space_pt"] = None
    return results

def get_style_font(style):
    """Get font properties from a style."""
    if style is None:
        return {}
    f = style.font
    return {
        "name": f.name,
        "size_pt": pt_val(f.size),
        "bold": f.bold,
        "italic": f.italic,
        "color": color_str(f.color),
    }

def get_style_paragraph(style):
    """Get paragraph format from a style."""
    if style is None:
        return {}
    pf = style.paragraph_format
    return {
        "alignment": ALIGN_MAP.get(pf.alignment, str(pf.alignment)),
        "first_line_indent_cm": emu_to_cm(pf.first_line_indent),
        "left_indent_cm": emu_to_cm(pf.left_indent),
        "right_indent_cm": emu_to_cm(pf.right_indent),
        "space_before_pt": emu_to_pt(pf.space_before),
        "space_after_pt": emu_to_pt(pf.space_after),
        "line_spacing": get_line_spacing(pf),
    }

def load_styles(doc):
    """Return dict of style_name -> style object for all paragraph styles."""
    result = {}
    for style in doc.styles:
        result[style.name] = style
    return result

def compare_dicts(d1, d2, label1="Template", label2="Generated"):
    """Return list of (key, val1, val2) where they differ."""
    diffs = []
    all_keys = set(list(d1.keys()) + list(d2.keys()))
    for k in sorted(all_keys):
        v1 = d1.get(k, "MISSING")
        v2 = d2.get(k, "MISSING")
        if str(v1) != str(v2):
            diffs.append((k, v1, v2))
    return diffs

def print_section(title):
    print("\n" + "="*80)
    print(f"  {title}")
    print("="*80)

def print_table(headers, rows):
    if not rows:
        print("  (no differences found)")
        return
    col_widths = [len(h) for h in headers]
    for row in rows:
        for i, cell in enumerate(row):
            col_widths[i] = max(col_widths[i], len(str(cell)))
    fmt = "  " + "  |  ".join(f"{{:<{w}}}" for w in col_widths)
    sep = "  " + "-+-".join("-" * w for w in col_widths)
    print(fmt.format(*headers))
    print(sep)
    for row in rows:
        print(fmt.format(*[str(c) for c in row]))

def get_run_info(run):
    f = run.font
    return {
        "text_preview": run.text[:30].replace('\n','\\n'),
        "name": f.name,
        "size_pt": pt_val(f.size),
        "bold": f.bold,
        "italic": f.italic,
        "color": color_str(f.color),
    }

def main():
    print("Loading documents...")
    tmpl = Document(TEMPLATE_PATH)
    gen = Document(GENERATED_PATH)

    tmpl_styles = load_styles(tmpl)
    gen_styles = load_styles(gen)

    # =========================================================================
    # 1. Page / Section settings
    # =========================================================================
    print_section("1. PAGE / SECTION SETTINGS")
    tmpl_secs = get_section_info(tmpl)
    gen_secs = get_section_info(gen)

    all_sec_keys = sorted(set(list(tmpl_secs.keys()) + list(gen_secs.keys())))
    for sk in all_sec_keys:
        t = tmpl_secs.get(sk, {})
        g = gen_secs.get(sk, {})
        diffs = compare_dicts(t, g)
        print(f"\n  [{sk}]")
        if diffs:
            print_table(["Property", "Template", "Generated"], diffs)
        else:
            print("  (identical)")

    # =========================================================================
    # 2. Style font formatting
    # =========================================================================
    print_section("2. STYLE FONT FORMATTING")
    for sname in STYLES_TO_CHECK:
        ts = tmpl_styles.get(sname)
        gs = gen_styles.get(sname)
        t_font = get_style_font(ts) if ts else {"status": "STYLE NOT FOUND"}
        g_font = get_style_font(gs) if gs else {"status": "STYLE NOT FOUND"}
        diffs = compare_dicts(t_font, g_font)
        if diffs or ts is None or gs is None:
            print(f"\n  Style: '{sname}'")
            if ts is None:
                print(f"    Template: STYLE NOT FOUND")
            if gs is None:
                print(f"    Generated: STYLE NOT FOUND")
            if diffs:
                print_table(["Font Property", "Template", "Generated"], diffs)

    # =========================================================================
    # 3. Style paragraph formatting
    # =========================================================================
    print_section("3. STYLE PARAGRAPH FORMATTING")
    for sname in STYLES_TO_CHECK:
        ts = tmpl_styles.get(sname)
        gs = gen_styles.get(sname)
        t_para = get_style_paragraph(ts) if ts else {"status": "STYLE NOT FOUND"}
        g_para = get_style_paragraph(gs) if gs else {"status": "STYLE NOT FOUND"}
        diffs = compare_dicts(t_para, g_para)
        if diffs or ts is None or gs is None:
            print(f"\n  Style: '{sname}'")
            if ts is None:
                print(f"    Template: STYLE NOT FOUND")
            if gs is None:
                print(f"    Generated: STYLE NOT FOUND")
            if diffs:
                print_table(["Para Property", "Template", "Generated"], diffs)

    # =========================================================================
    # 4. Run-level font overrides in first 5 non-empty paragraphs
    # =========================================================================
    print_section("4. RUN-LEVEL FONT OVERRIDES (first 5 non-empty paragraphs)")

    def analyze_doc_runs(doc, label):
        print(f"\n  --- {label} ---")
        count = 0
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            count += 1
            if count > 5:
                break
            style_name = para.style.name if para.style else "None"
            print(f"\n  Para #{count} (style='{style_name}'): \"{para.text[:60].replace(chr(10),' ')}\"")
            if not para.runs:
                print("    (no runs)")
                continue
            rows = []
            for j, run in enumerate(para.runs):
                if not run.text.strip():
                    continue
                ri = get_run_info(run)
                rows.append([
                    f"run[{j}]",
                    f"\"{ri['text_preview']}\"",
                    ri['name'],
                    ri['size_pt'],
                    ri['bold'],
                    ri['italic'],
                    ri['color'],
                ])
            if rows:
                print_table(["run", "text", "font.name", "size_pt", "bold", "italic", "color"], rows)
            else:
                print("    (all runs empty)")

    analyze_doc_runs(tmpl, "Template")
    analyze_doc_runs(gen, "Generated")

    # =========================================================================
    # 5. All styles in template vs generated
    # =========================================================================
    print_section("5. STYLE PRESENCE SUMMARY")
    tmpl_style_names = set(tmpl_styles.keys())
    gen_style_names = set(gen_styles.keys())
    only_in_tmpl = sorted(tmpl_style_names - gen_style_names)
    only_in_gen = sorted(gen_style_names - tmpl_style_names)
    print(f"\n  Styles only in Template ({len(only_in_tmpl)}): {only_in_tmpl[:30]}")
    print(f"\n  Styles only in Generated ({len(only_in_gen)}): {only_in_gen[:30]}")

    # =========================================================================
    # 6. Full comparison of all target styles (both font + para together)
    # =========================================================================
    print_section("6. FULL STYLE DETAILS (all target styles, template values)")
    for sname in STYLES_TO_CHECK:
        ts = tmpl_styles.get(sname)
        gs = gen_styles.get(sname)
        print(f"\n  Style: '{sname}'")
        print(f"    Template exists: {ts is not None}  |  Generated exists: {gs is not None}")
        if ts:
            tf = get_style_font(ts)
            tp = get_style_paragraph(ts)
            print(f"    [Template] Font: {tf}")
            print(f"    [Template] Para: {tp}")
        if gs:
            gf = get_style_font(gs)
            gp = get_style_paragraph(gs)
            print(f"    [Generated] Font: {gf}")
            print(f"    [Generated] Para: {gp}")

    print("\n\nDone.")

if __name__ == "__main__":
    main()
