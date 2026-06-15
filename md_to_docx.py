#!/usr/bin/env python3
"""Convert 算法课程论文.md to Word document."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

def set_font(run, name_cn='宋体', name_en='Times New Roman', size=None, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph_with_inlines(doc, text, style_name='Normal', align=None, font_size=11):
    """Parse inline bold (**...**) and code (`...`) in text, add to doc."""
    para = doc.add_paragraph(style=style_name)
    if align:
        para.alignment = align

    # Split by bold or inline code patterns
    pattern = re.compile(r'(\*\*.*?\*\*|`[^`]+`)')
    parts = pattern.split(text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, size=font_size, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            set_font(run, name_cn='Courier New', name_en='Courier New', size=font_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                run = para.add_run(part)
                set_font(run, size=font_size)
    return para

def add_code_block(doc, code_lines):
    """Add a shaded code block paragraph."""
    code_text = '\n'.join(code_lines)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    # Indent
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    pPr.append(ind)

    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)
    return para

def add_table(doc, rows):
    """rows: list of lists (first is header). Normalise row widths to header width."""
    if not rows:
        return
    col_count = len(rows[0])
    # Normalise: merge extra cells caused by pipe chars inside cell content
    normalised = []
    for row in rows:
        if len(row) <= col_count:
            # pad if short
            normalised.append(row + [''] * (col_count - len(row)))
        else:
            # merge surplus cells back into the last expected column
            merged = row[:col_count - 1]
            merged.append('|'.join(row[col_count - 1:]))
            normalised.append(merged)

    table = doc.add_table(rows=len(normalised), cols=col_count)
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(normalised):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=10, bold=(r_idx == 0))
                if r_idx == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def parse_table_row(line):
    line = line.strip().strip('|')
    return [c.strip() for c in line.split('|')]

def is_separator_row(row):
    return all(re.match(r'^[-:]+$', cell.strip()) for cell in row if cell.strip())

def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    i = 0
    in_code_block = False
    code_lang = ''
    code_lines = []
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block start/end
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:]
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith('|'):
            row = parse_table_row(line)
            if is_separator_row(row):
                i += 1
                continue
            table_rows.append(row)
            i += 1
            continue
        else:
            if table_rows:
                add_table(doc, table_rows)
                table_rows = []

        stripped = line.strip()

        # Heading 1 (Title)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            set_font(run, name_cn='黑体', name_en='Arial', size=16, bold=True)
            i += 1
            continue

        # Heading 2
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_font(run, name_cn='黑体', name_en='Arial', size=14, bold=True)
            i += 1
            continue

        # Heading 3
        if stripped.startswith('### '):
            text = stripped[4:]
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_font(run, name_cn='黑体', name_en='Arial', size=12, bold=True)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Metadata lines (bold key: value)
        meta_match = re.match(r'^\*\*(.+?)：\*\*\s*(.*)', stripped)
        if meta_match:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_key = para.add_run(meta_match.group(1) + '：')
            set_font(run_key, size=11, bold=True)
            run_val = para.add_run(meta_match.group(2))
            set_font(run_val, size=11)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2)
            para = add_paragraph_with_inlines(doc, text, font_size=11)
            pPr = para._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            ind.set(qn('w:hanging'), '360')
            pPr.append(ind)
            run0 = para.runs[0] if para.runs else para.add_run('')
            para.runs[0].text = f"{num_match.group(1)}. " + para.runs[0].text
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- '):
            text = stripped[2:]
            para = add_paragraph_with_inlines(doc, '• ' + text, font_size=11)
            pPr = para._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            pPr.append(ind)
            i += 1
            continue

        # Normal paragraph
        add_paragraph_with_inlines(doc, stripped, font_size=11)
        i += 1

    # Flush trailing table
    if table_rows:
        add_table(doc, table_rows)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    convert(
        '/Users/wen/Desktop/未命名文件夹/DeepLearning/算法课程论文.md',
        '/Users/wen/Desktop/未命名文件夹/DeepLearning/算法课程论文.docx'
    )
