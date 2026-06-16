"""
motor 译稿 docx 全面完善脚本
1. 插入 OMML 数学公式替换公式截图
2. 删除占位符段落
3. 清理公式说明冗余文字
4. 所有图注：拆分主标题与术语说明为两段
5. 修复列表项/特殊段落缩进
"""
import re, os, copy
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOC = "/Users/wen/Desktop/未命名文件夹/DeepLearning/522025320162_王文杰_分布式2026作业.docx"
W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M   = "http://schemas.openxmlformats.org/officeDocument/2006/math"

doc = Document(DOC)
body = doc.element.body

def wt(tag):  return f"{{{W}}}{tag}"
def mt(tag):  return f"{{{M}}}{tag}"

def remove_p_el(p_el):
    p_el.getparent().remove(p_el)

def add_run_to_p(p_el, text, bold=False, sz_half=None):
    r = etree.SubElement(p_el, wt("r"))
    rPr = etree.SubElement(r, wt("rPr"))
    if bold:
        etree.SubElement(rPr, wt("b"))
        etree.SubElement(rPr, wt("bCs"))
    if sz_half:
        for tag in (wt("sz"), wt("szCs")):
            el = etree.SubElement(rPr, tag)
            el.set(wt("val"), str(sz_half))
    t_el = etree.SubElement(r, wt("t"))
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t_el.text = text

def make_centered_body_p():
    """新建居中 Body Text 段落 XML 元素"""
    p = etree.Element(wt("p"))
    pPr = etree.SubElement(p, wt("pPr"))
    pStyle = etree.SubElement(pPr, wt("pStyle"))
    pStyle.set(wt("val"), "BodyText")
    jc = etree.SubElement(pPr, wt("jc"))
    jc.set(wt("val"), "center")
    ind = etree.SubElement(pPr, wt("ind"))
    ind.set(wt("firstLine"), "0")
    return p

# ──────────────────────────────────────────────────────────────────────────────
# 收集所有段落（带索引），方便后续查找
# ──────────────────────────────────────────────────────────────────────────────
all_paras = list(doc.paragraphs)  # 固定列表

def find_para(condition):
    """返回第一个满足条件的 (index, para) 或 (None, None)"""
    for i, p in enumerate(all_paras):
        if condition(i, p):
            return i, p
    return None, None

# ══════════════════════════════════════════════════════════════════════════════
# 1. 定位公式相关段落
# ══════════════════════════════════════════════════════════════════════════════
# 找"属性条带大小。"段落的索引
abs_size_idx, _ = find_para(lambda i, p: '属性条带大小' in p.text and '估算' in p.text)
print(f"  属性条带大小段落: idx={abs_size_idx}")

# 找公式图片段落（紧跟属性条带大小段落后，有图片的段落）
formula_img_idx = None
if abs_size_idx is not None:
    for i in range(abs_size_idx + 1, min(abs_size_idx + 4, len(all_paras))):
        p = all_paras[i]
        has_img = p._p.find('.//' + qn('w:drawing')) is not None
        if has_img:
            formula_img_idx = i
            break
print(f"  公式图片段落:     idx={formula_img_idx}")

# 找占位符段落
placeholder_idx, _ = find_para(lambda i, p: '公式：属性条带大小 ABS 估算' in p.text)
print(f"  占位符段落:       idx={placeholder_idx}")

# 找公式说明段落
explanation_idx, _ = find_para(lambda i, p: '其中 n 为不同 TotAttrSize' in p.text)
print(f"  公式说明段落:     idx={explanation_idx}")

# ══════════════════════════════════════════════════════════════════════════════
# 2. 构造 OMML 公式段落
#    ABS = Σ(i=1→n) max(VNum × Frequency_i, 1) × TotAttrSize_i
# ══════════════════════════════════════════════════════════════════════════════
FORMULA_XML = f"""<m:oMath
    xmlns:m="{M}"
    xmlns:w="{W}">
  <m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>ABS</m:t></m:r>
  <m:r><m:rPr/><m:t xml:space="preserve"> = </m:t></m:r>
  <m:nary>
    <m:naryPr>
      <m:chr m:val="&#x2211;"/>
      <m:limLoc m:val="undOvr"/>
      <m:grow/>
      <m:subHide m:val="0"/>
      <m:supHide m:val="0"/>
    </m:naryPr>
    <m:sub>
      <m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>i</m:t></m:r>
      <m:r><m:rPr/><m:t>=1</m:t></m:r>
    </m:sub>
    <m:sup>
      <m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>n</m:t></m:r>
    </m:sup>
    <m:e>
      <m:r><m:rPr><m:sty m:val="p"/></m:rPr>
        <m:t xml:space="preserve">max</m:t></m:r>
      <m:d>
        <m:dPr>
          <m:begChr m:val="("/>
          <m:endChr m:val=")"/>
        </m:dPr>
        <m:e>
          <m:r><m:rPr><m:sty m:val="i"/></m:rPr>
            <m:t xml:space="preserve">VNum</m:t></m:r>
          <m:r><m:rPr/><m:t xml:space="preserve"> \xd7 </m:t></m:r>
          <m:sSub>
            <m:sSubPr/>
            <m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr>
              <m:t>Frequency</m:t></m:r></m:e>
            <m:sub><m:r><m:rPr><m:sty m:val="i"/></m:rPr>
              <m:t>i</m:t></m:r></m:sub>
          </m:sSub>
          <m:r><m:rPr/><m:t xml:space="preserve">, 1</m:t></m:r>
        </m:e>
      </m:d>
      <m:r><m:rPr/><m:t xml:space="preserve"> \xd7 </m:t></m:r>
      <m:sSub>
        <m:sSubPr/>
        <m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr>
          <m:t>TotAttrSize</m:t></m:r></m:e>
        <m:sub><m:r><m:rPr><m:sty m:val="i"/></m:rPr>
          <m:t>i</m:t></m:r></m:sub>
      </m:sSub>
    </m:e>
  </m:nary>
</m:oMath>"""

omath_el = etree.fromstring(FORMULA_XML)

def make_omml_para():
    """MTDisplayEquation 样式的 OMML 段落"""
    p = etree.Element(wt("p"))
    pPr = etree.SubElement(p, wt("pPr"))
    pStyle = etree.SubElement(pPr, wt("pStyle"))
    pStyle.set(wt("val"), "MTDisplayEquation")
    jc = etree.SubElement(pPr, wt("jc"))
    jc.set(wt("val"), "center")
    oMathPara = etree.SubElement(p, mt("oMathPara"))
    oMathParaPr = etree.SubElement(oMathPara, mt("oMathParaPr"))
    jcm = etree.SubElement(oMathParaPr, mt("jc"))
    jcm.set(mt("val"), "center")
    oMathPara.append(copy.deepcopy(omath_el))
    return p

# 执行替换
if formula_img_idx is not None:
    img_p_el = all_paras[formula_img_idx]._p
    omml_p = make_omml_para()
    img_p_el.addprevious(omml_p)   # 在图片段落前插入公式
    remove_p_el(img_p_el)          # 删除图片段落
    print("✓ OMML 公式已插入（替换截图）")
else:
    print("⚠ 未找到公式截图段落，跳过")

# 重新加载（因为修改了 DOM，需要刷新段落列表）
doc.save(DOC)
doc = Document(DOC)
all_paras = list(doc.paragraphs)
body = doc.element.body

# ══════════════════════════════════════════════════════════════════════════════
# 3. 删除占位符，修复说明文字
# ══════════════════════════════════════════════════════════════════════════════
for p in list(doc.paragraphs):
    t = p.text.strip()
    if '公式：属性条带大小 ABS 估算' in t:
        remove_p_el(p._p)
        print("✓ 占位符已删除")
    elif '其中 n 为不同 TotAttrSize' in t:
        clean = t.replace("原文公式截图如下：", "").strip()
        # 清空原 run，重写
        for child in list(p._p):
            if child.tag not in (wt("pPr"),):
                p._p.remove(child)
        pPr = p._p.find(wt("pPr"))
        if pPr is None:
            pPr = etree.SubElement(p._p, wt("pPr"))
        jc = pPr.find(wt("jc"))
        if jc is None:
            jc = etree.SubElement(pPr, wt("jc"))
        jc.set(wt("val"), "center")
        ind = pPr.find(wt("ind"))
        if ind is None:
            ind = etree.SubElement(pPr, wt("ind"))
        ind.set(wt("firstLine"), "0")
        add_run_to_p(p._p, clean)
        print(f"✓ 公式说明已修复: {clean[:60]!r}")

# 保存后重载
doc.save(DOC)
doc = Document(DOC)
all_paras = list(doc.paragraphs)
body = doc.element.body

# ══════════════════════════════════════════════════════════════════════════════
# 4. 图注拆分：主标题 + 术语说明 → 两个段落
# ══════════════════════════════════════════════════════════════════════════════
split_count = 0
# 需要先收集要处理的段落，避免迭代中修改
targets = []
for p in all_paras:
    t = p.text.strip()
    if re.match(r'^图\d+[：:]', t) and '术语说明' in t:
        targets.append(p)
    # 图7和图8 合并行（含"图8"但无"术语说明"时也处理）
    elif re.match(r'^图7[：:]', t) and '图8' in t:
        targets.append(p)

for p in targets:
    t = p.text.strip()

    # 拆分策略
    if '术语说明' in t:
        parts = t.split('术语说明', 1)
        main_part = parts[0].rstrip('。').rstrip('，').strip()
        term_part = '术语说明' + parts[1].strip()
    elif '图8' in t and re.match(r'^图7', t):
        idx8 = t.index('图8')
        main_part = t[:idx8].rstrip()
        term_part = t[idx8:]
    else:
        continue

    # 修改原段落 → 只含主标题
    for child in list(p._p):
        if child.tag != wt("pPr"):
            p._p.remove(child)

    # 写入主标题（图号加粗）
    m_obj = re.match(r'^(图\d+[：:]+\s*)', main_part)
    if m_obj:
        add_run_to_p(p._p, m_obj.group(1), bold=True)
        rest = main_part[m_obj.end():]
        if rest:
            add_run_to_p(p._p, rest, bold=False)
    else:
        add_run_to_p(p._p, main_part)

    # 在原段落后插入术语说明段落（8 pt，居中）
    term_p_el = make_centered_body_p()
    add_run_to_p(term_p_el, term_part, bold=False, sz_half=16)  # 8 pt
    p._p.addnext(term_p_el)
    split_count += 1

print(f"✓ 图注拆分 {split_count} 处")

# 重载确保 DOM 一致
doc.save(DOC)
doc = Document(DOC)

# ══════════════════════════════════════════════════════════════════════════════
# 5. 修复列表项缩进
# ══════════════════════════════════════════════════════════════════════════════
list_fixed = 0
for p in doc.paragraphs:
    if p.style.name == 'Body Text' and p.text.strip().startswith('•'):
        pPr = p._p.find(wt("pPr"))
        if pPr is None:
            pPr = etree.SubElement(p._p, wt("pPr"))
        ind = pPr.find(wt("ind"))
        if ind is None:
            ind = etree.SubElement(pPr, wt("ind"))
        ind.set(wt("left"),      "480")
        ind.set(wt("firstLine"), "0")
        list_fixed += 1
print(f"✓ 列表项缩进 {list_fixed} 处")

# ══════════════════════════════════════════════════════════════════════════════
# 6. 通讯作者/原文出处首行缩进清零
# ══════════════════════════════════════════════════════════════════════════════
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('*Corresponding') or t.startswith('原文：OSDI'):
        pPr = p._p.find(wt("pPr"))
        if pPr is None:
            pPr = etree.SubElement(p._p, wt("pPr"))
        ind = pPr.find(wt("ind"))
        if ind is None:
            ind = etree.SubElement(pPr, wt("ind"))
        ind.set(wt("firstLine"), "0")
print("✓ 特殊段落首行缩进清零")

# ══════════════════════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════════════════════
doc.save(DOC)
sz = os.path.getsize(DOC)
plen = len(doc.paragraphs)
print(f"\n✓ 完成：{DOC}")
print(f"  大小 {sz//1024} KB，共 {plen} 段落")
