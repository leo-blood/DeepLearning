#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Deep comparison: actual paragraph content vs style definitions,
checking for run-level overrides and XML-level font settings.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE_PATH = "/Users/wen/Desktop/未命名文件夹/DeepLearning/论文/软件学报2016年排版样例模版.docx"
GENERATED_PATH = "/Users/wen/Desktop/未命名文件夹/DeepLearning/motor_osdi24_zh.docx"

STYLES_TO_CHECK = [
    "Subtitle", "作者", "Name", "摘要", "Abstract",
    "关键词", "Key words", "Heading 1", "Heading 2",
    "Body Text", "Text of Reference", "Reference"
]

def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 3)

def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / 12700, 2)

def halfpt_to_pt(val):
    """Convert half-point to pt."""
    if val is None:
        return None
    return round(int(val) / 2, 1)

def twip_to_pt(val):
    """Convert twip (1/20 pt) to pt."""
    if val is None:
        return None
    return round(int(val) / 20, 2)

def twip_to_cm(val):
    """Convert twip to cm."""
    if val is None:
        return None
    return round(int(val) / 20 / 72 * 2.54, 3)

def get_xml_font_info(el):
    """Extract font info from an XML element's rPr."""
    rPr = el.find(qn('w:rPr'))
    if rPr is None:
        return {}
    result = {}

    # Font name
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        result['rFonts_ascii'] = rFonts.get(qn('w:ascii'))
        result['rFonts_eastAsia'] = rFonts.get(qn('w:eastAsia'))
        result['rFonts_hAnsi'] = rFonts.get(qn('w:hAnsi'))
        result['rFonts_cs'] = rFonts.get(qn('w:cs'))

    # Font size
    sz = rPr.find(qn('w:sz'))
    szCs = rPr.find(qn('w:szCs'))
    if sz is not None:
        result['sz_pt'] = halfpt_to_pt(sz.get(qn('w:val')))
    if szCs is not None:
        result['szCs_pt'] = halfpt_to_pt(szCs.get(qn('w:val')))

    # Bold
    b = rPr.find(qn('w:b'))
    bCs = rPr.find(qn('w:bCs'))
    if b is not None:
        val = b.get(qn('w:val'))
        result['bold'] = False if val == '0' else True
    if bCs is not None:
        val = bCs.get(qn('w:val'))
        result['boldCs'] = False if val == '0' else True

    # Color
    color = rPr.find(qn('w:color'))
    if color is not None:
        result['color'] = color.get(qn('w:val'))

    # Lang
    lang = rPr.find(qn('w:lang'))
    if lang is not None:
        result['lang_val'] = lang.get(qn('w:val'))
        result['lang_eastAsia'] = lang.get(qn('w:eastAsia'))

    return result

def get_xml_para_info(pPr):
    """Extract paragraph format from pPr XML element."""
    if pPr is None:
        return {}
    result = {}

    # Justification
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        result['jc'] = jc.get(qn('w:val'))

    # Indentation
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        result['ind_left_twip'] = ind.get(qn('w:left'))
        result['ind_right_twip'] = ind.get(qn('w:right'))
        result['ind_firstLine_twip'] = ind.get(qn('w:firstLine'))
        result['ind_hanging_twip'] = ind.get(qn('w:hanging'))
        if result['ind_left_twip']:
            result['ind_left_cm'] = twip_to_cm(result['ind_left_twip'])
        if result['ind_firstLine_twip']:
            result['ind_firstLine_cm'] = twip_to_cm(result['ind_firstLine_twip'])
        if result['ind_hanging_twip']:
            result['ind_hanging_cm'] = twip_to_cm(result['ind_hanging_twip'])

    # Spacing
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        result['spacing_before_twip'] = spacing.get(qn('w:before'))
        result['spacing_after_twip'] = spacing.get(qn('w:after'))
        result['spacing_line_twip'] = spacing.get(qn('w:line'))
        result['spacing_lineRule'] = spacing.get(qn('w:lineRule'))
        if result['spacing_before_twip']:
            result['spacing_before_pt'] = twip_to_pt(result['spacing_before_twip'])
        if result['spacing_after_twip']:
            result['spacing_after_pt'] = twip_to_pt(result['spacing_after_twip'])
        if result['spacing_line_twip']:
            result['spacing_line_pt'] = twip_to_pt(result['spacing_line_twip'])

    return result

def get_style_xml_full(style):
    """Get complete XML-level style info."""
    if style is None:
        return {}
    result = {}

    # Style XML element
    el = style.element

    # rPr at style level
    rPr = el.find(qn('w:rPr'))
    if rPr is not None:
        result['style_rPr'] = get_xml_font_info(el)

    # pPr at style level
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        result['style_pPr'] = get_xml_para_info(pPr)

    return result

def analyze_paragraphs_by_style(doc, style_name, max_paras=3):
    """Find paragraphs using a given style and analyze their XML."""
    results = []
    count = 0
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name:
            if not para.text.strip():
                continue
            count += 1
            if count > max_paras:
                break

            para_info = {
                'text_preview': para.text[:50].replace('\n', '\\n'),
                'para_pPr': get_xml_para_info(para._p.find(qn('w:pPr'))),
                'runs': []
            }

            for run in para.runs:
                if not run.text.strip():
                    continue
                r_el = run._r
                run_info = {
                    'text': run.text[:20],
                    'rPr': get_xml_font_info(r_el)
                }
                para_info['runs'].append(run_info)

            results.append(para_info)

    return results

def print_section(title):
    print("\n" + "="*80)
    print(f"  {title}")
    print("="*80)

def compare_style_xml(tmpl_doc, gen_doc, style_name):
    """Compare actual paragraph XML for a given style."""
    print(f"\n  Style: '{style_name}'")

    t_paras = analyze_paragraphs_by_style(tmpl_doc, style_name, 2)
    g_paras = analyze_paragraphs_by_style(gen_doc, style_name, 2)

    if not t_paras and not g_paras:
        print("    (no paragraphs found with this style in either doc)")
        return

    # Compare pPr overrides
    if t_paras and g_paras:
        t_pPr = t_paras[0]['para_pPr']
        g_pPr = g_paras[0]['para_pPr']

        all_keys = set(list(t_pPr.keys()) + list(g_pPr.keys()))
        pPr_diffs = []
        for k in sorted(all_keys):
            tv = t_pPr.get(k, 'MISSING')
            gv = g_pPr.get(k, 'MISSING')
            if str(tv) != str(gv):
                pPr_diffs.append((k, tv, gv))

        if pPr_diffs:
            print(f"    Para-level pPr DIFFERENCES:")
            for k, tv, gv in pPr_diffs:
                print(f"      {k}: Template={tv}  |  Generated={gv}")
        else:
            print(f"    Para-level pPr: identical (or both empty)")
            if t_pPr:
                print(f"    Para pPr values: {t_pPr}")

    # Compare run rPr overrides
    print(f"    Run-level rPr comparison:")

    for label, paras in [("Template", t_paras), ("Generated", g_paras)]:
        if not paras:
            print(f"      {label}: no paragraphs found")
            continue
        para = paras[0]
        print(f"      {label} para: \"{para['text_preview']}\"")
        if not para['runs']:
            print(f"        (no runs with content)")
        else:
            run = para['runs'][0]
            rPr = run['rPr']
            if not rPr:
                print(f"        run[0] \"{run['text']}\": NO rPr overrides (inherits from style)")
            else:
                print(f"        run[0] \"{run['text']}\": rPr overrides = {rPr}")

def check_section_xml(doc, label):
    """Deep XML check of section properties."""
    print(f"\n  [{label}]")
    for i, section in enumerate(doc.sections):
        sectPr = section._sectPr
        print(f"  section[{i}] XML:")

        # Page size
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            w = pgSz.get(qn('w:w'))
            h = pgSz.get(qn('w:h'))
            orient = pgSz.get(qn('w:orient'))
            print(f"    pgSz: w={w} ({twip_to_cm(w) if w else None}cm), h={h} ({twip_to_cm(h) if h else None}cm), orient={orient}")

        # Page margins
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            top = pgMar.get(qn('w:top'))
            bottom = pgMar.get(qn('w:bottom'))
            left = pgMar.get(qn('w:left'))
            right = pgMar.get(qn('w:right'))
            header = pgMar.get(qn('w:header'))
            footer = pgMar.get(qn('w:footer'))
            print(f"    pgMar: top={top}({twip_to_cm(top) if top else None}cm), bottom={bottom}({twip_to_cm(bottom) if bottom else None}cm)")
            print(f"           left={left}({twip_to_cm(left) if left else None}cm), right={right}({twip_to_cm(right) if right else None}cm)")
            print(f"           header={header}({twip_to_pt(header) if header else None}pt), footer={footer}({twip_to_pt(footer) if footer else None}pt)")

        # Columns
        cols = sectPr.find(qn('w:cols'))
        if cols is not None:
            num = cols.get(qn('w:num'))
            space = cols.get(qn('w:space'))
            equalWidth = cols.get(qn('w:equalWidth'))
            print(f"    cols: num={num}, space={space}({twip_to_cm(space) if space else None}cm), equalWidth={equalWidth}")
            # Individual column definitions
            col_els = cols.findall(qn('w:col'))
            for ci, col_el in enumerate(col_els):
                cw = col_el.get(qn('w:w'))
                csp = col_el.get(qn('w:space'))
                print(f"      col[{ci}]: w={cw}({twip_to_cm(cw) if cw else None}cm), space={csp}({twip_to_cm(csp) if csp else None}cm)")
        else:
            print(f"    cols: (not set, single column)")

def check_style_xml_deep(doc, style_name, label):
    """Print full XML of a style's rPr and pPr."""
    style = None
    for s in doc.styles:
        if s.name == style_name:
            style = s
            break

    if style is None:
        print(f"  [{label}] Style '{style_name}': NOT FOUND")
        return

    el = style.element
    rPr = el.find(qn('w:rPr'))
    pPr = el.find(qn('w:pPr'))

    print(f"  [{label}] Style '{style_name}':")
    if rPr is not None:
        print(f"    rPr XML: {etree.tostring(rPr, pretty_print=False).decode()}")
    else:
        print(f"    rPr: (none)")
    if pPr is not None:
        print(f"    pPr XML: {etree.tostring(pPr, pretty_print=False).decode()}")
    else:
        print(f"    pPr: (none)")

def get_default_font(doc):
    """Get document default fonts."""
    # Check docDefaults in styles.xml
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        return "No docDefaults"

    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    pPrDefault = docDefaults.find(qn('w:pPrDefault'))

    result = {}
    if rPrDefault is not None:
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                result['default_ascii'] = rFonts.get(qn('w:ascii'))
                result['default_eastAsia'] = rFonts.get(qn('w:eastAsia'))
                result['default_hAnsi'] = rFonts.get(qn('w:hAnsi'))
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                result['default_sz_pt'] = halfpt_to_pt(sz.get(qn('w:val')))
            lang = rPr.find(qn('w:lang'))
            if lang is not None:
                result['default_lang'] = lang.get(qn('w:val'))
                result['default_lang_eastAsia'] = lang.get(qn('w:eastAsia'))

    return result

def find_all_paragraphs_by_style(doc, style_name):
    """Find all paragraphs with a given style and show their text."""
    results = []
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name:
            results.append(para.text[:80])
    return results

def main():
    print("Loading documents...")
    tmpl = Document(TEMPLATE_PATH)
    gen = Document(GENERATED_PATH)

    # =========================================================================
    # Section XML deep dive
    # =========================================================================
    print_section("1. SECTION XML DEEP DIVE")
    check_section_xml(tmpl, "Template")
    check_section_xml(gen, "Generated")

    # =========================================================================
    # Document default fonts
    # =========================================================================
    print_section("2. DOCUMENT DEFAULT FONTS")
    t_defaults = get_default_font(tmpl)
    g_defaults = get_default_font(gen)
    print(f"  Template defaults: {t_defaults}")
    print(f"  Generated defaults: {g_defaults}")

    # Differences
    all_keys = set(list(t_defaults.keys()) + list(g_defaults.keys()))
    diffs = [(k, t_defaults.get(k,'MISSING'), g_defaults.get(k,'MISSING'))
             for k in sorted(all_keys)
             if str(t_defaults.get(k,'MISSING')) != str(g_defaults.get(k,'MISSING'))]
    if diffs:
        print(f"\n  DIFFERENCES:")
        for k, tv, gv in diffs:
            print(f"    {k}: Template={tv}  |  Generated={gv}")
    else:
        print(f"\n  (no differences in document defaults)")

    # =========================================================================
    # Style XML deep dive for key styles
    # =========================================================================
    print_section("3. STYLE XML DEEP DIVE (key styles)")
    key_styles = ["Subtitle", "作者", "Body Text", "Heading 1", "Abstract", "摘要"]
    for sname in key_styles:
        print(f"\n  --- '{sname}' ---")
        check_style_xml_deep(tmpl, sname, "Template")
        check_style_xml_deep(gen, sname, "Generated")

    # =========================================================================
    # Actual paragraph pPr/rPr overrides for each style
    # =========================================================================
    print_section("4. ACTUAL PARAGRAPH OVERRIDES BY STYLE")
    for sname in STYLES_TO_CHECK:
        compare_style_xml(tmpl, gen, sname)

    # =========================================================================
    # Check theme fonts (word/theme/theme1.xml)
    # =========================================================================
    print_section("5. THEME FONTS")
    for label, doc in [("Template", tmpl), ("Generated", gen)]:
        try:
            theme_part = doc.part.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme')
            theme_xml = theme_part._blob
            from lxml import etree as ET
            tree = ET.fromstring(theme_xml)
            ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            fontScheme = tree.find(f'.//{{{ns}}}fontScheme')
            if fontScheme is not None:
                majorFont = fontScheme.find(f'{{{ns}}}majorFont')
                minorFont = fontScheme.find(f'{{{ns}}}minorFont')
                print(f"  [{label}]")
                if majorFont is not None:
                    latin = majorFont.find(f'{{{ns}}}latin')
                    ea = majorFont.find(f'{{{ns}}}ea')
                    cs = majorFont.find(f'{{{ns}}}cs')
                    print(f"    majorFont: latin={latin.get('typeface') if latin is not None else None}, ea={ea.get('typeface') if ea is not None else None}, cs={cs.get('typeface') if cs is not None else None}")
                if minorFont is not None:
                    latin = minorFont.find(f'{{{ns}}}latin')
                    ea = minorFont.find(f'{{{ns}}}ea')
                    cs = minorFont.find(f'{{{ns}}}cs')
                    print(f"    minorFont: latin={latin.get('typeface') if latin is not None else None}, ea={ea.get('typeface') if ea is not None else None}, cs={cs.get('typeface') if cs is not None else None}")
        except Exception as e:
            print(f"  [{label}] Could not read theme: {e}")

    # =========================================================================
    # Check Subtitle style paragraphs - what's actually in them
    # =========================================================================
    print_section("6. SUBTITLE PARAGRAPH FULL XML")
    for label, doc in [("Template", tmpl), ("Generated", gen)]:
        print(f"\n  [{label}]")
        for para in doc.paragraphs:
            if para.style and para.style.name == "Subtitle" and para.text.strip():
                p_el = para._p
                print(f"  Full XML: {etree.tostring(p_el, pretty_print=True).decode()[:3000]}")
                break

    # =========================================================================
    # Check Body Text paragraph full XML in generated
    # =========================================================================
    print_section("7. BODY TEXT PARAGRAPH FULL XML (Generated, first 3)")
    count = 0
    for para in gen.paragraphs:
        if para.style and para.style.name == "Body Text" and para.text.strip():
            count += 1
            if count > 3:
                break
            p_el = para._p
            print(f"\n  Para #{count}: \"{para.text[:60]}\"")
            print(f"  XML: {etree.tostring(p_el, pretty_print=False).decode()[:1000]}")

    # =========================================================================
    # Check Heading 1 paragraphs
    # =========================================================================
    print_section("8. HEADING 1 PARAGRAPH FULL XML (both docs, first 2)")
    for label, doc in [("Template", tmpl), ("Generated", gen)]:
        print(f"\n  [{label}]")
        count = 0
        for para in doc.paragraphs:
            if para.style and para.style.name in ("Heading 1", "标题1") and para.text.strip():
                count += 1
                if count > 2:
                    break
                p_el = para._p
                print(f"  Para #{count}: \"{para.text[:60]}\"")
                print(f"  XML: {etree.tostring(p_el, pretty_print=False).decode()[:800]}")

    # =========================================================================
    # List all unique styles used in generated doc
    # =========================================================================
    print_section("9. ALL STYLES USED IN GENERATED DOC")
    used_styles = {}
    for para in gen.paragraphs:
        sname = para.style.name if para.style else "None"
        if sname not in used_styles:
            used_styles[sname] = 0
        used_styles[sname] += 1

    print("\n  Style Name                          | Count")
    print("  " + "-"*50)
    for sname, cnt in sorted(used_styles.items(), key=lambda x: -x[1]):
        print(f"  {sname:<35} | {cnt}")

    print("\n\nDone.")

if __name__ == "__main__":
    main()
