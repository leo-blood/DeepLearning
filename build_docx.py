"""
Motor 论文中译版 docx 生成脚本
依据软件学报模版样式生成 motor_osdi24_zh.docx
"""
import os, re
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE    = "/Users/wen/Desktop/未命名文件夹/DeepLearning"
TEMPLATE= os.path.join(BASE, "论文/软件学报2016年排版样例模版.docx")
FIGS    = os.path.join(BASE, "figures")
MD_FILE = os.path.join(BASE, "motor_osdi24_zh.md")
OUTPUT  = os.path.join(BASE, "motor_osdi24_zh.docx")

# ─────────────────────────────────────────────
# 打开模版并清空正文
# ─────────────────────────────────────────────
doc = Document(TEMPLATE)
body = doc.element.body
for child in list(body):
    if not child.tag.endswith("}sectPr"):
        body.remove(child)

# ─────────────────────────────────────────────
# 辅助函数
# ─────────────────────────────────────────────
def clean(text):
    """处理 md 转义字符，把 \* 还原为 *，去除行内 $...$"""
    text = text.replace("\\*", "*")
    # 把行内数学 $...$ 简单去掉美元符号保留内容
    text = re.sub(r'\$([^$]+?)\$', lambda m: m.group(1), text)
    return text

def add_runs(para, text):
    """把包含 **bold** 的文本逐 run 加入段落"""
    text = clean(text)
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (idx % 2 == 1)

def new_para(style, text="", center=False, indent_zero=False):
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if center or indent_zero:
        p.paragraph_format.first_line_indent = 0
    if text:
        add_runs(p, text)
    return p

def insert_image(rel_path, width=Inches(5.0)):
    full = os.path.join(BASE, rel_path) if not os.path.isabs(rel_path) else rel_path
    if not os.path.exists(full):
        return
    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run().add_picture(full, width=width)

def insert_fig_caption(text):
    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    add_runs(p, text)

# ─────────────────────────────────────────────
# 读取 md 文件
# ─────────────────────────────────────────────
with open(MD_FILE, encoding="utf-8") as f:
    lines = f.readlines()

# ─────────────────────────────────────────────
# 页眉信息：中文标题
# ─────────────────────────────────────────────
new_para("Subtitle", "Motor：在分离式内存上为分布式事务启用多版本机制")

# ─────────────────────────────────────────────
# 作者 & 单位
# ─────────────────────────────────────────────
new_para("作者", "Ming Zhang, Yu Hua*, Zhijun Yang")
new_para("Name",
         "Wuhan National Laboratory for Optoelectronics, School of Computer, "
         "Huazhong University of Science and Technology")

# 通讯作者 + 原文出处
p = doc.add_paragraph(style="Body Text")
p.paragraph_format.first_line_indent = 0
p.add_run("*Corresponding Author: Yu Hua (csyhua@hust.edu.cn)")

p2 = doc.add_paragraph(style="Body Text")
p2.paragraph_format.first_line_indent = 0
p2.add_run("原文：OSDI 2024, 第 801–819 页")

# ─────────────────────────────────────────────
# 中文摘要 + 关键词
# ─────────────────────────────────────────────
ZH_ABSTRACT = (
    "在现代数据中心中，内存分离技术将单体服务器解耦，构建出通过网络互连的独立计算池和内存"
    "池，从而提升资源利用率并提供高性能服务。计算池通过分布式事务访问内存池中的远程数据，以"
    "保证原子性和强一致性。现有的单版本设计由于并发度有限、日志开销高而受到制约。虽然传统单"
    "体服务器中的多版本设计有望提供高并发并降低日志开销，但其无法直接应用于分离式内存。为了"
    "弥合多版本设计与分离式内存之间的鸿沟，我们提出了 Motor——通过整体性地重新设计版本结构"
    "和事务协议，在分离式内存上实现快速分布式事务处理的多版本机制。为了在内存池中高效组织同"
    "一数据的不同版本，Motor 采用一种新颖的连续版本元组（CVT）结构，将各版本连续存储于连续"
    "地址空间中，使计算池只需一次网络往返即可获取目标版本。在 CVT 之上，Motor 设计了一套完"
    "全基于单侧 RDMA 的 MVCC 协议，支持具有灵活隔离级别的快速分布式事务。实验结果表明，与"
    "现有最优系统相比，Motor 将吞吐量最高提升 98.1%，将延迟最高降低 55.8%。"
)
p = doc.add_paragraph(style="摘要")
r = p.add_run("摘  要：\t")
add_runs(p, ZH_ABSTRACT)

p = doc.add_paragraph(style="关键词")
p.add_run("关键词：\t内存分离；多版本并发控制；分布式事务；RDMA；连续版本元组")

# ─────────────────────────────────────────────
# 英文摘要 + 关键词
# ─────────────────────────────────────────────
EN_ABSTRACT = (
    "In modern datacenters, memory disaggregation unpacks monolithic servers to build "
    "network-connected distributed compute and memory pools to improve resource utilization "
    "and deliver high performance. The compute pool leverages distributed transactions to "
    "access remote data in the memory pool to provide atomicity and strong consistency. "
    "Existing single-versioning designs have been constrained due to limited system "
    "concurrency and high logging overheads. Although the multi-versioning design in the "
    "conventional monolithic servers is promising to offer high concurrency and reduce "
    "logging overheads, which however fails to work in the disaggregated memory. In order "
    "to bridge the gap between the multi-versioning design and the disaggregated memory, "
    "we propose Motor that holistically redesigns the version structure and transaction "
    "protocol to enable multi-versioning for fast distributed transaction processing on the "
    "disaggregated memory. To efficiently organize different versions of data in the memory "
    "pool, Motor leverages a new consecutive version tuple (CVT) structure to store the "
    "versions together in a continuous manner, which allows the compute pool to obtain the "
    "target version in a single network round trip. On top of CVT, Motor leverages a fully "
    "one-sided RDMA-based MVCC protocol to support fast distributed transactions with "
    "flexible isolation levels. Experimental results demonstrate that Motor improves the "
    "throughput by up to 98.1% and reduces the latency by up to 55.8% compared with "
    "state-of-the-art systems."
)
p = doc.add_paragraph(style="Abstract")
r = p.add_run("Abstract:  ")
r.bold = True
p.add_run(EN_ABSTRACT)

p = doc.add_paragraph(style="Key words")
r = p.add_run("Key words: ")
r.bold = True
p.add_run("disaggregated memory; multi-version concurrency control (MVCC); "
          "distributed transactions; RDMA; consecutive version tuple (CVT)")

# ─────────────────────────────────────────────
# 逐行解析 md 正文（从 ## 1. 引言 开始）
# ─────────────────────────────────────────────
# 定位正文起始行（第一个 ## 1. 开头的行）
start_line = 0
for idx, ln in enumerate(lines):
    if re.match(r'^## 1\.', ln.strip()):
        start_line = idx
        break

i = start_line
total = len(lines)

# 用于判断是否进入参考文献区
in_refs = False

def is_fig_caption_line(s):
    return bool(re.match(r'\*\*图\d', s))

while i < total:
    raw = lines[i].rstrip("\n")
    s   = raw.strip()
    i  += 1

    if not s:
        continue

    # ── 分隔线 ──
    if s == "---":
        continue

    # ── 图片行 ──
    m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', s)
    if m:
        rel  = m.group(2)
        # 跳过 formula_abs.png，下面用 $$ 块处理
        if "formula_abs" in rel:
            continue
        insert_image(rel, width=Inches(5.2))
        continue

    # ── 块公式行（单行 $$...$$ 或 开头 $$）──
    if s.startswith("$$"):
        formula_img = os.path.join(FIGS, "formula_abs.png")
        if os.path.exists(formula_img):
            insert_image(formula_img, width=Inches(4.8))
            # 公式说明文字
            p = doc.add_paragraph(style="Body Text")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = 0
            p.add_run("（公式：属性条带大小 ABS 估算，详见正文第 4.2 节）")
        # 若是多行块，跳到结束 $$
        if s == "$$":
            while i < total and lines[i].strip() != "$$":
                i += 1
            i += 1  # 跳过结束 $$
        continue

    # ── 图注行（以 **图N 开头）──
    if is_fig_caption_line(s):
        insert_fig_caption(s)
        continue

    # ── 参考文献标题 ──
    if re.match(r'^## 参考文献', s):
        in_refs = True
        new_para("Reference", "参考文献")
        continue

    # ── 参考文献条目 ──
    if in_refs:
        if re.match(r'^\[\d+\]', s):
            new_para("Text of Reference", s)
        else:
            new_para("Text of Reference", s)
        continue

    # ── 各级标题 ──
    if re.match(r'^## \d+\.', s):
        # 一级节标题：## 1. 引言
        text = re.sub(r'^## \d+\.\s*', '', s)
        new_para("Heading 1", text)
        continue

    if re.match(r'^##\s', s):
        text = re.sub(r'^##\s+', '', s)
        new_para("Heading 1", text)
        continue

    if re.match(r'^### \d+\.\d+\s', s):
        text = re.sub(r'^### \d+\.\d+\s*', '', s)
        new_para("Heading 2", text)
        continue

    if re.match(r'^###\s', s):
        text = re.sub(r'^###\s+', '', s)
        new_para("Heading 2", text)
        continue

    # ── 列表项 ──
    if s.startswith("- "):
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.first_line_indent = 0
        p.paragraph_format.left_indent = Cm(0.5)
        add_runs(p, "• " + s[2:])
        continue

    # ── 跳过 **Abstract** 独占行 ──
    if s == "**Abstract**":
        continue

    # ── 跳过 ## 摘要 行（已在前面单独处理）──
    if s in ("## 摘要", "**Abstract**"):
        continue

    # ── 普通正文 ──
    new_para("Body Text", s)

# ─────────────────────────────────────────────
# 保存
# ─────────────────────────────────────────────
doc.save(OUTPUT)
print(f"✓ 已生成：{OUTPUT}")
print(f"  文件大小：{os.path.getsize(OUTPUT)/1024:.0f} KB")
