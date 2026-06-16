#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Check all paragraph-level pPr overrides in generated doc for each target style.
Focus on: Body Text firstLine, and run-level bold=False override issue.
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE_PATH = "/Users/wen/Desktop/未命名文件夹/DeepLearning/论文/软件学报2016年排版样例模版.docx"
GENERATED_PATH = "/Users/wen/Desktop/未命名文件夹/DeepLearning/motor_osdi24_zh.docx"

STYLES_TO_CHECK = [
    "Subtitle", "作者", "Name", "摘要", "Abstract",
    "关键词", "Key words", "Heading 1", "Heading 2",
    "Body Text", "Text of Reference", "Reference"
]

def twip_to_cm(val):
    if val is None:
        return None
    return round(int(val) / 20 / 72 * 2.54, 3)

def twip_to_pt(val):
    if val is None:
        return None
    return round(int(val) / 20, 2)

def halfpt_to_pt(val):
    if val is None:
        return None
    return round(int(val) / 2, 1)

def get_pPr_overrides(para):
    """Get paragraph-level pPr overrides (not from style)."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return {}
    result = {}

    # Indentation
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        left = ind.get(qn('w:left'))
        right = ind.get(qn('w:right'))
        firstLine = ind.get(qn('w:firstLine'))
        hanging = ind.get(qn('w:hanging'))
        firstLineChars = ind.get(qn('w:firstLineChars'))
        hangingChars = ind.get(qn('w:hangingChars'))
        leftChars = ind.get(qn('w:leftChars'))
        if left: result['left_cm'] = twip_to_cm(left)
        if right: result['right_cm'] = twip_to_cm(right)
        if firstLine is not None: result['firstLine_cm'] = twip_to_cm(firstLine)
        if hanging: result['hanging_cm'] = twip_to_cm(hanging)
        if firstLineChars is not None: result['firstLineChars'] = firstLineChars
        if hangingChars: result['hangingChars'] = hangingChars
        if leftChars: result['leftChars'] = leftChars

    # Spacing
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        before = spacing.get(qn('w:before'))
        after = spacing.get(qn('w:after'))
        line = spacing.get(qn('w:line'))
        lineRule = spacing.get(qn('w:lineRule'))
        if before is not None: result['space_before_pt'] = twip_to_pt(before)
        if after is not None: result['space_after_pt'] = twip_to_pt(after)
        if line is not None: result['line_pt'] = twip_to_pt(line)
        if lineRule: result['lineRule'] = lineRule

    # Justification
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        result['jc'] = jc.get(qn('w:val'))

    return result

def get_run_overrides(para):
    """Get run-level rPr overrides for all runs in a paragraph."""
    overrides = []
    for run in para.runs:
        r_el = run._r
        rPr = r_el.find(qn('w:rPr'))
        if rPr is None:
            continue

        run_info = {'text': run.text[:25]}

        # Bold
        b = rPr.find(qn('w:b'))
        if b is not None:
            val = b.get(qn('w:val'))
            run_info['bold'] = 'False (suppressed)' if val == '0' else 'True'

        # Font name
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            ascii_f = rFonts.get(qn('w:ascii'))
            ea_f = rFonts.get(qn('w:eastAsia'))
            if ascii_f: run_info['font_ascii'] = ascii_f
            if ea_f: run_info['font_eastAsia'] = ea_f

        # Size
        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            run_info['sz_pt'] = halfpt_to_pt(sz.get(qn('w:val')))

        # Color
        color = rPr.find(qn('w:color'))
        if color is not None:
            run_info['color'] = color.get(qn('w:val'))

        if len(run_info) > 1:  # has overrides beyond text
            overrides.append(run_info)

    return overrides

def print_section(title):
    print("\n" + "="*80)
    print(f"  {title}")
    print("="*80)

def main():
    gen = Document(GENERATED_PATH)
    tmpl = Document(TEMPLATE_PATH)

    # =========================================================================
    # 1. Body Text paragraphs: check firstLine override in generated
    # =========================================================================
    print_section("1. BODY TEXT: firstLine override in generated (first 10 paragraphs)")
    count = 0
    for para in gen.paragraphs:
        if para.style and para.style.name == "Body Text" and para.text.strip():
            count += 1
            if count > 10:
                break
            ppr = get_pPr_overrides(para)
            run_ov = get_run_overrides(para)
            text_preview = para.text[:50].replace('\n', '\\n')
            print(f"\n  Para #{count}: \"{text_preview}\"")
            print(f"    pPr overrides: {ppr if ppr else '(none)'}")
            if run_ov:
                for r in run_ov[:3]:
                    print(f"    run override: {r}")

    # =========================================================================
    # 2. All styles: count of paragraphs with bold=False run override
    # =========================================================================
    print_section("2. BOLD=FALSE RUN OVERRIDE STATISTICS (Generated)")
    style_bold_counts = {}
    for para in gen.paragraphs:
        sname = para.style.name if para.style else "None"
        has_bold_false = False
        for run in para.runs:
            r_el = run._r
            rPr = r_el.find(qn('w:rPr'))
            if rPr is not None:
                b = rPr.find(qn('w:b'))
                if b is not None:
                    val = b.get(qn('w:val'))
                    if val == '0':
                        has_bold_false = True
                        break
        if has_bold_false:
            style_bold_counts[sname] = style_bold_counts.get(sname, 0) + 1

    print(f"\n  Paragraphs with bold=False (w:b val='0') run override:")
    for sname, cnt in sorted(style_bold_counts.items(), key=lambda x: -x[1]):
        print(f"    {sname:<35} | {cnt} paragraphs")

    # =========================================================================
    # 3. Template: same analysis
    # =========================================================================
    print_section("3. BOLD=FALSE RUN OVERRIDE STATISTICS (Template)")
    tmpl_bold_counts = {}
    for para in tmpl.paragraphs:
        sname = para.style.name if para.style else "None"
        has_bold_false = False
        for run in para.runs:
            r_el = run._r
            rPr = r_el.find(qn('w:rPr'))
            if rPr is not None:
                b = rPr.find(qn('w:b'))
                if b is not None:
                    val = b.get(qn('w:val'))
                    if val == '0':
                        has_bold_false = True
                        break
        if has_bold_false:
            tmpl_bold_counts[sname] = tmpl_bold_counts.get(sname, 0) + 1

    print(f"\n  Paragraphs with bold=False (w:b val='0') run override:")
    if tmpl_bold_counts:
        for sname, cnt in sorted(tmpl_bold_counts.items(), key=lambda x: -x[1]):
            print(f"    {sname:<35} | {cnt} paragraphs")
    else:
        print("    (none)")

    # =========================================================================
    # 4. Heading 1 paragraphs in generated: full analysis
    # =========================================================================
    print_section("4. HEADING 1 PARAGRAPHS IN GENERATED: FULL ANALYSIS")
    for para in gen.paragraphs:
        if para.style and para.style.name == "Heading 1" and para.text.strip():
            ppr = get_pPr_overrides(para)
            run_ov = get_run_overrides(para)
            print(f"\n  \"{para.text[:60]}\"")
            print(f"    pPr: {ppr if ppr else '(none)'}")
            if run_ov:
                for r in run_ov[:5]:
                    print(f"    run: {r}")

    # =========================================================================
    # 5. Template Heading 1 paragraphs: full analysis
    # =========================================================================
    print_section("5. HEADING 1 PARAGRAPHS IN TEMPLATE: FULL ANALYSIS")
    for para in tmpl.paragraphs:
        if para.style and para.style.name in ("Heading 1", "标题1") and para.text.strip():
            ppr = get_pPr_overrides(para)
            run_ov = get_run_overrides(para)
            print(f"\n  \"{para.text[:60]}\"")
            print(f"    pPr: {ppr if ppr else '(none)'}")
            if run_ov:
                for r in run_ov[:5]:
                    print(f"    run: {r}")

    # =========================================================================
    # 6. Abstract paragraph in both docs
    # =========================================================================
    print_section("6. ABSTRACT PARAGRAPHS: FULL ANALYSIS")
    for label, doc in [("Template", tmpl), ("Generated", gen)]:
        print(f"\n  [{label}]")
        for para in doc.paragraphs:
            if para.style and para.style.name == "Abstract" and para.text.strip():
                ppr = get_pPr_overrides(para)
                run_ov = get_run_overrides(para)
                print(f"  \"{para.text[:70]}\"")
                print(f"    pPr overrides: {ppr if ppr else '(none)'}")
                if run_ov:
                    for r in run_ov[:5]:
                        print(f"    run: {r}")
                break

    # =========================================================================
    # 7. Key words paragraph in both docs
    # =========================================================================
    print_section("7. KEY WORDS PARAGRAPH: FULL ANALYSIS")
    for label, doc in [("Template", tmpl), ("Generated", gen)]:
        print(f"\n  [{label}]")
        for para in doc.paragraphs:
            if para.style and para.style.name == "Key words" and para.text.strip():
                ppr = get_pPr_overrides(para)
                run_ov = get_run_overrides(para)
                print(f"  \"{para.text[:70]}\"")
                print(f"    pPr overrides: {ppr if ppr else '(none)'}")
                if run_ov:
                    for r in run_ov[:5]:
                        print(f"    run: {r}")
                break

    # =========================================================================
    # 8. Heading 2 paragraphs in generated
    # =========================================================================
    print_section("8. HEADING 2 PARAGRAPHS IN GENERATED (first 5)")
    count = 0
    for para in gen.paragraphs:
        if para.style and para.style.name == "Heading 2" and para.text.strip():
            count += 1
            if count > 5:
                break
            ppr = get_pPr_overrides(para)
            run_ov = get_run_overrides(para)
            print(f"\n  #{count} \"{para.text[:60]}\"")
            print(f"    pPr: {ppr if ppr else '(none)'}")
            if run_ov:
                for r in run_ov[:3]:
                    print(f"    run: {r}")

    # =========================================================================
    # 9. Template Heading 2 paragraphs
    # =========================================================================
    print_section("9. HEADING 2 PARAGRAPHS IN TEMPLATE (first 5)")
    count = 0
    for para in tmpl.paragraphs:
        if para.style and para.style.name == "Heading 2" and para.text.strip():
            count += 1
            if count > 5:
                break
            ppr = get_pPr_overrides(para)
            run_ov = get_run_overrides(para)
            print(f"\n  #{count} \"{para.text[:60]}\"")
            print(f"    pPr: {ppr if ppr else '(none)'}")
            if run_ov:
                for r in run_ov[:3]:
                    print(f"    run: {r}")

    # =========================================================================
    # 10. Text of Reference in generated (first 5)
    # =========================================================================
    print_section("10. TEXT OF REFERENCE IN GENERATED (first 5)")
    count = 0
    for para in gen.paragraphs:
        if para.style and para.style.name == "Text of Reference" and para.text.strip():
            count += 1
            if count > 5:
                break
            ppr = get_pPr_overrides(para)
            run_ov = get_run_overrides(para)
            print(f"\n  #{count} \"{para.text[:60]}\"")
            print(f"    pPr: {ppr if ppr else '(none)'}")
            if run_ov:
                for r in run_ov[:3]:
                    print(f"    run: {r}")

    print("\n\nDone.")

if __name__ == "__main__":
    main()
