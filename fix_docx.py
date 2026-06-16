"""
修复 motor_osdi24_zh.docx 与模版的格式差异
主要修复项：
  P0  所有 run 携带 bold=False 覆盖 → 批量删除
  P1  图注/通讯作者等段落 firstLine=0 → 删除错误覆盖
  P2  Abstract 缺 space_before / line_spacing 覆盖
  P2  Heading 2 缺 space_before / space_after 覆盖
  P3  Key words / 关键词 悬挂缩进覆盖缺失
  P3  Name 段落缺 space_before / space_after 覆盖
  P3  Subtitle 补 run 字号覆盖（14 pt，与模版一致）
"""
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import copy, os

SRC = "/Users/wen/Desktop/未命名文件夹/DeepLearning/motor_osdi24_zh.docx"
OUT = SRC   # 原地修复

doc = Document(SRC)

# ──────────────────────────────────────────────────────────
# P0  删除所有 run 中的 <w:b w:val="0"/> 覆盖
#     （以及 <w:bCs w:val="0"/>，中文加粗同理）
# ──────────────────────────────────────────────────────────
removed_b = 0
for para in doc.paragraphs:
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is None:
            continue
        for tag in (qn("w:b"), qn("w:bCs")):
            el = rpr.find(tag)
            if el is not None:
                val = el.get(qn("w:val"), "1")
                # val="0" 或 val="false" 表示强制关闭
                if val in ("0", "false"):
                    rpr.remove(el)
                    removed_b += 1

print(f"P0  已删除 bold=False 覆盖：{removed_b} 处")

# ──────────────────────────────────────────────────────────
# P1  Body Text 段落中被错误设置 firstLine=0 的情况
#     （图注、通讯作者等段落需要 firstLine=0 + center，保留）
#     （正文普通段落：firstLine 不应为 0，直接删除 pPr 中的 ind 覆盖）
# ──────────────────────────────────────────────────────────
fixed_indent = 0
for para in doc.paragraphs:
    if para.style.name != "Body Text":
        continue
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        continue   # 图注/公式行保持 center + firstLine=0
    ppr = para._p.find(qn("w:pPr"))
    if ppr is None:
        continue
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        continue
    first = ind.get(qn("w:firstLine"), None)
    if first == "0":
        # 删除整个 ind 元素，让段落继承样式的首行缩进
        ppr.remove(ind)
        fixed_indent += 1

print(f"P1  已修复 Body Text firstLine=0 覆盖：{fixed_indent} 处")

# ──────────────────────────────────────────────────────────
# 辅助：给段落 pPr 写入 spacing 覆盖
# ──────────────────────────────────────────────────────────
def set_spacing(para, before_twip=None, after_twip=None,
                line_twip=None, line_rule="exact"):
    ppr = para._p.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.SubElement(para._p, qn("w:pPr"))
    spc = ppr.find(qn("w:spacing"))
    if spc is None:
        spc = etree.SubElement(ppr, qn("w:spacing"))
    if before_twip is not None:
        spc.set(qn("w:before"), str(before_twip))
    if after_twip is not None:
        spc.set(qn("w:after"), str(after_twip))
    if line_twip is not None:
        spc.set(qn("w:line"), str(line_twip))
        rule_map = {"exact": "exact", "auto": "auto", "atLeast": "atLeast"}
        spc.set(qn("w:lineRule"), rule_map.get(line_rule, "exact"))

def set_indent(para, left_twip=None, hanging_twip=None, first_line_twip=None):
    ppr = para._p.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.SubElement(para._p, qn("w:pPr"))
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = etree.SubElement(ppr, qn("w:ind"))
    if left_twip is not None:
        ind.set(qn("w:left"), str(left_twip))
    if hanging_twip is not None:
        ind.set(qn("w:hanging"), str(hanging_twip))
    if first_line_twip is not None:
        ind.set(qn("w:firstLine"), str(first_line_twip))

def set_run_sz(run, sz_half_pt):
    """设置 run 字号（sz 单位是 half-point）"""
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run._r, qn("w:rPr"))
    for tag in (qn("w:sz"), qn("w:szCs")):
        el = rpr.find(tag)
        if el is None:
            el = etree.SubElement(rpr, tag)
        el.set(qn("w:val"), str(sz_half_pt))

# ──────────────────────────────────────────────────────────
# P2  Abstract 段落：space_before=142 twip，line=240 twip(exact)
# ──────────────────────────────────────────────────────────
for para in doc.paragraphs:
    if para.style.name == "Abstract":
        set_spacing(para, before_twip=142, line_twip=240, line_rule="exact")

print("P2  Abstract 段落 spacing 已修复")

# ──────────────────────────────────────────────────────────
# P2  Heading 2 段落：space_before=71 twip，space_after=71 twip
# ──────────────────────────────────────────────────────────
cnt = 0
for para in doc.paragraphs:
    if para.style.name == "Heading 2":
        set_spacing(para, before_twip=71, after_twip=71)
        cnt += 1

print(f"P2  Heading 2 spacing 已修复：{cnt} 个段落")

# ──────────────────────────────────────────────────────────
# P3  Key words 段落：left=1979 twip, hanging=1979 twip
#     (1.746 cm ≈ 990 twip；模版实测约 990 twip)
# ──────────────────────────────────────────────────────────
KW_TWIP = 990   # 1.746 cm ≈ 990 twip
for para in doc.paragraphs:
    if para.style.name == "Key words":
        set_indent(para, left_twip=KW_TWIP, hanging_twip=KW_TWIP)

print("P3  Key words 悬挂缩进已修复")

# ──────────────────────────────────────────────────────────
# P3  关键词 段落：left=800 twip, hanging=800 twip
#     (1.408 cm ≈ 800 twip)
# ──────────────────────────────────────────────────────────
KW_ZH_TWIP = 800
for para in doc.paragraphs:
    if para.style.name == "关键词":
        set_indent(para, left_twip=KW_ZH_TWIP, hanging_twip=KW_ZH_TWIP)

print("P3  关键词 悬挂缩进已修复")

# ──────────────────────────────────────────────────────────
# P3  Name 段落：space_before=100 twip，space_after=100 twip
#     (5pt ≈ 100 twip)
# ──────────────────────────────────────────────────────────
for para in doc.paragraphs:
    if para.style.name == "Name":
        set_spacing(para, before_twip=100, after_twip=100)

print("P3  Name spacing 已修复")

# ──────────────────────────────────────────────────────────
# P3  Subtitle 段落：run 字号设为 14 pt（与模版实测一致）
#     spacing before=0, line=atLeast 0
# ──────────────────────────────────────────────────────────
for para in doc.paragraphs:
    if para.style.name == "Subtitle":
        set_spacing(para, before_twip=0, line_twip=0, line_rule="atLeast")
        for run in para.runs:
            set_run_sz(run, sz_half_pt=280)   # 14 pt = 280 half-pt

print("P3  Subtitle 字号/行距已修复")

# ──────────────────────────────────────────────────────────
# 额外：摘要 段落行距 — 模版为 240 twip exact（9pt字 + 固定行距）
# ──────────────────────────────────────────────────────────
for para in doc.paragraphs:
    if para.style.name == "摘要":
        set_spacing(para, line_twip=240, line_rule="exact")

print("补充 摘要 行距已修复")

# ──────────────────────────────────────────────────────────
# 额外：关键词 段落行距同样修复
# ──────────────────────────────────────────────────────────
for para in doc.paragraphs:
    if para.style.name == "关键词":
        set_spacing(para, line_twip=240, line_rule="exact")

# ──────────────────────────────────────────────────────────
# 保存
# ──────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\n✓ 修复完成，已保存：{OUT}")
print(f"  文件大小：{os.path.getsize(OUT)/1024:.0f} KB")
